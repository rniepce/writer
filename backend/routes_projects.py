from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime
import io
import re
import json

from database import get_db
from models import Project, Chapter, Scene, Character, ReferenceDoc

router = APIRouter(prefix="/projects", tags=["projects"])


# Pydantic Schemas
class ProjectCreate(BaseModel):
    title: str
    description: Optional[str] = None


class ProjectUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None


class ProjectResponse(BaseModel):
    id: int
    title: str
    description: Optional[str]
    created_at: datetime
    chapter_count: int

    class Config:
        from_attributes = True


class ProjectCard(BaseModel):
    """Minimal data for project list"""
    id: int
    title: str
    chapter_count: int
    created_at: datetime

    class Config:
        from_attributes = True


# Routes
@router.get("", response_model=List[ProjectCard])
def list_projects(db: Session = Depends(get_db)):
    """List all projects with chapter counts"""
    projects = db.query(Project).order_by(Project.created_at.desc()).all()
    return [
        ProjectCard(
            id=p.id,
            title=p.title,
            chapter_count=len(p.chapters),
            created_at=p.created_at
        )
        for p in projects
    ]


@router.get("/{project_id}/export")
def export_project(project_id: int, format: str = "docx", db: Session = Depends(get_db)):
    """Export all chapters of a project as .docx or .pdf"""
    if format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Formato inválido. Use 'docx' ou 'pdf'.")

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Projeto não encontrado")

    chapters = (
        db.query(Chapter)
        .filter(Chapter.project_id == project_id)
        .order_by(Chapter.order)
        .all()
    )

    safe_title = re.sub(r'[^\w\s-]', '', project.title).strip().replace(' ', '_')

    if format == "docx":
        return _export_docx(safe_title, project.title, chapters)
    else:
        return _export_pdf(safe_title, project.title, chapters)


@router.get("/{project_id}/save")
def save_project_offline(project_id: int, db: Session = Depends(get_db)):
    """Export the entire project as a .writerproject JSON bundle for offline backup / portability."""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Projeto não encontrado")

    chapters = (
        db.query(Chapter)
        .filter(Chapter.project_id == project_id)
        .order_by(Chapter.order)
        .all()
    )

    chapters_data = []
    for ch in chapters:
        scenes = (
            db.query(Scene)
            .filter(Scene.chapter_id == ch.id)
            .order_by(Scene.order)
            .all()
        )
        chapters_data.append({
            "title": ch.title,
            "content": ch.content or "",
            "order": ch.order,
            "word_count": ch.word_count,
            "color": ch.color,
            "scenes": [
                {"title": s.title, "content": s.content or "", "order": s.order}
                for s in scenes
            ],
        })

    characters = db.query(Character).filter(Character.project_id == project_id).all()
    characters_data = [
        {"name": c.name, "description": c.description or ""}
        for c in characters
    ]

    ref_docs = db.query(ReferenceDoc).filter(ReferenceDoc.project_id == project_id).all()
    ref_docs_data = [
        {"doc_type": r.doc_type, "content": r.content or "", "filename": r.filename}
        for r in ref_docs
    ]

    bundle = {
        "writer_version": "1.0",
        "exported_at": datetime.utcnow().isoformat(),
        "project": {
            "title": project.title,
            "description": project.description or "",
            "chapters": chapters_data,
            "characters": characters_data,
            "reference_docs": ref_docs_data,
        },
    }

    safe_title = re.sub(r'[^\w\s-]', '', project.title).strip().replace(' ', '_')
    json_bytes = json.dumps(bundle, ensure_ascii=False, indent=2).encode("utf-8")
    buf = io.BytesIO(json_bytes)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.writerproject"'},
    )


@router.post("/import", response_model=ProjectResponse)
async def import_project(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Import a .writerproject JSON bundle and recreate the project with all data."""
    try:
        raw = await file.read()
        bundle = json.loads(raw)
    except (json.JSONDecodeError, UnicodeDecodeError):
        raise HTTPException(status_code=400, detail="Arquivo inválido. Envie um .writerproject válido.")

    if "writer_version" not in bundle or "project" not in bundle:
        raise HTTPException(status_code=400, detail="Formato de arquivo não reconhecido.")

    proj_data = bundle["project"]

    # Create project
    db_project = Project(
        title=proj_data.get("title", "Projeto Importado"),
        description=proj_data.get("description"),
    )
    db.add(db_project)
    db.flush()  # get the id

    # Chapters + Scenes
    for ch_data in proj_data.get("chapters", []):
        db_chapter = Chapter(
            project_id=db_project.id,
            title=ch_data.get("title", "Capítulo"),
            content=ch_data.get("content", ""),
            order=ch_data.get("order", 0),
            word_count=ch_data.get("word_count", 0),
            color=ch_data.get("color"),
        )
        db.add(db_chapter)
        db.flush()

        for sc_data in ch_data.get("scenes", []):
            db_scene = Scene(
                chapter_id=db_chapter.id,
                title=sc_data.get("title", "Cena"),
                content=sc_data.get("content", ""),
                order=sc_data.get("order", 0),
            )
            db.add(db_scene)

    # Characters
    for char_data in proj_data.get("characters", []):
        db_char = Character(
            project_id=db_project.id,
            name=char_data.get("name", "Personagem"),
            description=char_data.get("description", ""),
        )
        db.add(db_char)

    # Reference docs
    for ref_data in proj_data.get("reference_docs", []):
        doc_type = ref_data.get("doc_type")
        if doc_type:
            db_ref = ReferenceDoc(
                project_id=db_project.id,
                doc_type=doc_type,
                content=ref_data.get("content", ""),
                filename=ref_data.get("filename"),
            )
            db.add(db_ref)

    db.commit()
    db.refresh(db_project)

    return ProjectResponse(
        id=db_project.id,
        title=db_project.title,
        description=db_project.description,
        created_at=db_project.created_at,
        chapter_count=len(db_project.chapters),
    )


def _html_to_plain(html: str) -> str:
    """Strip HTML tags and return plain text."""
    from bs4 import BeautifulSoup
    if not html:
        return ""
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n").strip()


def _export_docx(safe_title: str, project_title: str, chapters):
    """Generate a .docx file from chapters."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup

    doc = Document()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(project_title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    doc.add_page_break()

    # Chapters
    for chapter in chapters:
        doc.add_heading(chapter.title, level=1)

        if chapter.content:
            soup = BeautifulSoup(chapter.content, "html.parser")
            for element in soup.children:
                tag = getattr(element, 'name', None)
                text = element.get_text(separator=" ").strip() if hasattr(element, 'get_text') else str(element).strip()
                if not text:
                    continue

                if tag in ('h1', 'h2', 'h3'):
                    level = int(tag[1]) + 1  # offset since chapter title is h1
                    doc.add_heading(text, level=min(level, 4))
                elif tag == 'blockquote':
                    para = doc.add_paragraph(text)
                    para.style = 'Quote'
                elif tag in ('ul', 'ol'):
                    for li in element.find_all('li', recursive=False):
                        doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                else:
                    # Regular paragraph – preserve bold / italic inline
                    para = doc.add_paragraph()
                    _add_inline_runs(para, element)
        else:
            doc.add_paragraph("(capítulo vazio)")

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )


def _add_inline_runs(para, element):
    """Walk through inline children and add runs preserving bold/italic."""
    from bs4 import NavigableString, Tag

    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            para.add_run(text)
        return

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip() or text == ' ':
                para.add_run(text)
        elif isinstance(child, Tag):
            text = child.get_text()
            if not text:
                continue
            run = para.add_run(text)
            if child.name in ('strong', 'b'):
                run.bold = True
            elif child.name in ('em', 'i'):
                run.italic = True
            elif child.name == 'u':
                run.underline = True


def _export_pdf(safe_title: str, project_title: str, chapters):
    """Generate a PDF file from chapters using weasyprint."""
    import weasyprint

    html_parts = [
        "<html><head><meta charset='utf-8'>",
        "<style>",
        "body { font-family: Georgia, 'Times New Roman', serif; max-width: 700px; margin: 0 auto; padding: 40px; color: #1a1a1a; line-height: 1.8; }",
        "h1 { text-align: center; font-size: 2em; margin-top: 200px; margin-bottom: 10px; }",
        ".chapter-title { page-break-before: always; font-size: 1.6em; text-align: center; margin-top: 60px; margin-bottom: 30px; border-bottom: 1px solid #ccc; padding-bottom: 10px; }",
        "p { text-indent: 2em; margin: 0.5em 0; text-align: justify; }",
        "blockquote { border-left: 3px solid #ccc; padding-left: 1em; margin: 1em 0; font-style: italic; color: #555; }",
        "</style>",
        "</head><body>",
        f"<h1>{project_title}</h1>",
        "<div style='page-break-after: always;'></div>",
    ]

    for chapter in chapters:
        html_parts.append(f"<h2 class='chapter-title'>{chapter.title}</h2>")
        if chapter.content:
            html_parts.append(chapter.content)
        else:
            html_parts.append("<p><em>(capítulo vazio)</em></p>")

    html_parts.append("</body></html>")
    full_html = "\n".join(html_parts)

    pdf_bytes = weasyprint.HTML(string=full_html).write_pdf()
    buf = io.BytesIO(pdf_bytes)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
    )


@router.get("/{project_id}", response_model=ProjectResponse)
def get_project(project_id: int, db: Session = Depends(get_db)):
    """Get project details"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Projeto não encontrado")
    return ProjectResponse(
        id=project.id,
        title=project.title,
        description=project.description,
        created_at=project.created_at,
        chapter_count=len(project.chapters)
    )


@router.post("", response_model=ProjectResponse)
def create_project(project: ProjectCreate, db: Session = Depends(get_db)):
    """Create a new project"""
    db_project = Project(
        title=project.title,
        description=project.description
    )
    db.add(db_project)
    db.commit()
    db.refresh(db_project)
    return ProjectResponse(
        id=db_project.id,
        title=db_project.title,
        description=db_project.description,
        created_at=db_project.created_at,
        chapter_count=0
    )


@router.put("/{project_id}", response_model=ProjectResponse)
def update_project(project_id: int, update: ProjectUpdate, db: Session = Depends(get_db)):
    """Update project details"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Projeto não encontrado")
    
    if update.title is not None:
        project.title = update.title
    if update.description is not None:
        project.description = update.description
    
    db.commit()
    db.refresh(project)
    return ProjectResponse(
        id=project.id,
        title=project.title,
        description=project.description,
        created_at=project.created_at,
        chapter_count=len(project.chapters)
    )


@router.delete("/{project_id}")
def delete_project(project_id: int, db: Session = Depends(get_db)):
    """Delete a project and all its chapters"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Projeto não encontrado")
    
    db.delete(project)
    db.commit()
    return {"message": "Projeto removido"}
